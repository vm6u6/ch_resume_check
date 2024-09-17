from docx import Document
from docx.shared import Inches

class output_newLayer():
    def __init__(self):
        self.document = Document()

    def output_newLayer(self, modified_section):

        self.document.add_heading(modified_section["姓名"], 0)

        for section, content in modified_section.items():
            self.document.add_paragraph(section)
            self.document.add_paragraph(content)    
            